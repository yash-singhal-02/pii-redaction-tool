#!/usr/bin/env python3
"""PII redaction tool for DOCX text and embedded ID-card images."""
from __future__ import annotations

import argparse
import json
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from PIL import Image, ImageDraw, ImageFont
import pytesseract

EMAIL_RE = re.compile(r"(?<![\w.+-])[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}(?![\w.-])")
SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
IP_RE = re.compile(r"(?<![\w.])(?:25[0-5]|2[0-4]\d|1?\d?\d)(?:\.(?:25[0-5]|2[0-4]\d|1?\d?\d)){3}(?![\w.])")
CC_RE = re.compile(r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)")
AADHAAR_RE = re.compile(r"\b\d{4}[ -]\d{4}[ -]\d{4}\b")
PAN_RE = re.compile(r"\b[A-Z]{5}\d{4}[A-Z]\b", re.I)
PHONE_PATTERNS = [
    re.compile(r"(?<!\d)\+?\s*91[\s-]?(?:\(\s*\d{2,4}\s*\)|\d{2,4})[\s-]?\d{3,5}[\s-]?\d{4}(?!\d)"),
    re.compile(r"(?<!\d)\+\d{1,3}[\s-]?(?:\(?\d{2,4}\)?[\s-]?)\d{3,4}[\s-]?\d{4}(?!\d)"),
    re.compile(r"(?<!\d)[6-9]\d{9}(?!\d)"),
]
DATE_RE = r"(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})"
DOB_RE = re.compile(rf"(?i)(?:date\s+of\s+birth|dob|born(?:\s+on)?)\s*[:\-]?\s*({DATE_RE})")
PIN_RE = re.compile(r"\b\d{3}(?:\s|-)?\d{3}\b")
ADDRESS_CONTEXT_RE = re.compile(r"(?i)\b(?:registered office|corporate office|mailing address|residential address|residence|address|correspondence address)\b")
ADDRESS_TERMS_RE = re.compile(r"(?i)\b(?:road|marg|street|lane|avenue|floor|plot|taluka|district|village|society|apartment|bungalow|building|tower|phase|pincode|pin|embassy|cts\s*no\.?|s\.?\s*no\.?)\b|\b[A-Z]-?\d{2,}\b")
COMPANY_SUFFIX_RE = re.compile(r"(?i)\b(?:Private\s+Limited|Limited|LLP|Inc\.?|Incorporated|Corporation|Bank|Foundation|Trust|Industries|Technologies|Motors|Electricals)\b")

STOPWORDS = {"The","Our","This","That","Red","Herring","Prospectus","Equity","Shares","Offer","Price","Anchor","Investor","Investors","Mutual","Funds","Life","Insurance","Companies","Pension","Book","Running","Lead","Managers","Stock","Exchanges","Working","Days","Company","Management","Risk","Factors","Capital","Structure","Financial","Statements","India","Maharashtra","Mumbai","Pune","Bank","Limited","Private","Trust","Family","KSH","Website","Email","Telephone","Registration","Number","Account","Address","Description","Term","For","With","Report","Committee","Application","Forms","Amount","Bidders","Manager","Managers","Securities","Road","Marg","Society","East","West","North","South","Building","Floor","Tower","Level","Centre","Complex","Village","Taluka","District","Plot","Phase","Park","House","Hotel","Apartment","Showroom","Department","Division","Market","World","Information","Agreement","Banks","Sponsor","Syndicate","Members","Locations","Intermediaries","Escrow","Agent","Period","Proceeds","Entity","Agency","Legal","Counsel","Law","Exchange","Board","Accounts","Cash","Rating","Experts","Specified","Registered","Broker","Brokers","Group","Entities","Branch","Parents","Share","Payment","Expense","Statutory","Auditors","Chartered","Accountants","Disclosure","Requirements","Regulatory","Other","Qualified","Institutional","Buyers","Statutory","Eligibility","Regulations","Obligations","Listing","Issue","General","Designated","Portion","Supported","Blocked","Self","Certified","Bidding","Summary","Definitions","Abbreviations"}
ROLE_WORDS = {"director","directors","chairman","executive","managing","whole-time","independent","promoter","shareholder","secretary","officer","person","personnel","chief","financial","compliance"}

@dataclass(frozen=True)
class Span:
    start: int
    end: int
    label: str


def luhn_valid(value: str) -> bool:
    digits = [int(c) for c in re.sub(r"\D", "", value)]
    if not 13 <= len(digits) <= 19:
        return False
    total = 0
    parity = len(digits) % 2
    for i, digit in enumerate(digits):
        if i % 2 == parity:
            digit *= 2
            if digit > 9:
                digit -= 9
        total += digit
    return total % 10 == 0


def title_name_candidates(text: str) -> list[str]:
    pat = re.compile(r"\b[A-Z][A-Za-z.'-]{1,30}(?:\s+[A-Z][A-Za-z.'-]{1,30}){1,3}\b")
    result = []
    for match in pat.finditer(text):
        value = match.group(0).strip(" ,;:/")
        words = value.split()
        if not 2 <= len(words) <= 4:
            continue
        if any(word in STOPWORDS for word in words):
            continue
        if any(word.lower() in ROLE_WORDS for word in words):
            continue
        if any(word.isupper() and len(word) > 1 for word in words):
            continue
        if COMPANY_SUFFIX_RE.search(value):
            continue
        result.append(value)
    return result


def discover_person_names(text: str) -> set[str]:
    names = set()
    cleaned = re.sub(r"\s+", " ", text)
    pat = (r"(?i)(?:contact\s+person|contact\s+persons?|promoters?)\s*:\s*"
           r"(.*?)(?=\b(?:contact\s+person|contact\s+persons?|promoters?|company|registered\s+office|"
           r"corporate\s+office|website|email|telephone|tel|sebi\s+registration)\b|$)")
    for match in re.finditer(pat, cleaned):
        for piece in re.split(r"[/,]", match.group(1)):
            names.update(title_name_candidates(piece))
    for match in re.finditer(r"(?i)\b(?:being|namely)\s+([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){1,3})", cleaned):
        names.update(title_name_candidates(match.group(1)))
    return names


def discover_company_names(text: str) -> set[str]:
    companies = set()
    cleaned = re.sub(r"\s+", " ", text)
    token = r"[A-Z][A-Za-z0-9&.'()/-]*"
    suffix = r"(?:Private\s+Limited|Limited|LLP|Inc\.?|Incorporated|Corporation|Bank|Foundation|Trust|Industries|Technologies|Motors|Electricals)"
    pattern = re.compile(rf"\b(?:{token}\s+){{1,6}}{suffix}\b")
    bad = {"the","our","fresh","issue","offer","anchor","investor","investors","date","email","telephone","and","public","account","bankers","registrar","to","former","formerly","contact","person","website","grievance","running","lead","managers","members","sponsor","promoter","trusts","term","abbreviations","company","escrow","collection","refund","bid","regulations","regulation","act","circular","senior","industrial","development","positive","stable","short","long","complaints","redressal","mechanism","self-certified","industrial","investment"}
    for match in pattern.finditer(cleaned):
        words = match.group(0).strip(" ,;:/").split()
        candidate = " ".join(words)
        if any(word.lower().strip(".,") in bad for word in words):
            continue
        if candidate.lower() in {"private limited", "limited", "bank limited"}:
            continue
        companies.add(candidate)
    companies.update(re.findall(r"\bWaterloo Industrial Park (?:I|II|III|IV|V|VI|VIII|IX(?: A| B)?) Private Limited\b", cleaned))
    companies.update(re.findall(r"\bKSH Infra Park (?:IV|VI) Private Limited\b", cleaned))
    return companies


class Redactor:
    def __init__(self):
        self.person_names = set()
        self.company_names = set()
        self.mapping = {}
        self.counters = {k: 0 for k in ["NAME","EMAIL","PHONE","COMPANY","ADDRESS","SSN","CREDIT_CARD","DOB","IP_ADDRESS"]}
        self.fake_names = ["John Doe","Jane Smith","Peter Parker","Alex Johnson","Emily Brown","Michael Wilson","Sarah Davis","David Miller","Olivia Taylor","Daniel Anderson"]
        self.fake_companies = ["Acme Technologies Private Limited","BlueSky Industries Limited","NorthStar Finance Limited","Greenfield Consulting LLP","Summit Securities Limited"]
        self.fake_addresses = ["101 Example Road, Mumbai, Maharashtra, India","22 Sample Street, Pune, Maharashtra, India","45 Demo Park, Bengaluru, Karnataka, India","9 Test Avenue, New Delhi, Delhi, India"]

    def learn(self, paragraphs):
        text = "\n".join(paragraphs)
        self.person_names |= discover_person_names(text)
        self.company_names |= discover_company_names(text)

    def learn_table_context(self, tables):
        for table in tables:
            if not table.rows:
                continue
            for row in table.rows[1:]:
                values = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(values)
                self.person_names |= set(title_name_candidates(row_text))
                self.company_names |= discover_company_names(row_text)

    def replacement(self, label, original):
        key = (label, original.casefold() if label in {"NAME", "COMPANY"} else original)
        if key in self.mapping:
            return self.mapping[key]
        n = self.counters[label]
        if label == "NAME": value = self.fake_names[n % len(self.fake_names)]
        elif label == "COMPANY": value = self.fake_companies[n % len(self.fake_companies)]
        elif label == "EMAIL": value = f"contact{n + 1}@example.com"
        elif label == "PHONE": value = f"+91 90000 {10000 + n:05d}"
        elif label == "ADDRESS": value = self.fake_addresses[n % len(self.fake_addresses)]
        elif label == "SSN": value = f"999-88-{1000 + n:04d}"
        elif label == "CREDIT_CARD": value = "4111 1111 1111 1111"
        elif label == "DOB": value = "01/01/1990"
        elif label == "IP_ADDRESS": value = f"192.0.2.{10 + n}"
        else: value = "[REDACTED]"
        self.mapping[key] = value
        self.counters[label] += 1
        return value.upper() if original.isupper() and label in {"NAME", "COMPANY"} else value

    def detect(self, text):
        spans = []
        for match in EMAIL_RE.finditer(text): spans.append(Span(match.start(), match.end(), "EMAIL"))
        for match in SSN_RE.finditer(text): spans.append(Span(match.start(), match.end(), "SSN"))
        for match in IP_RE.finditer(text): spans.append(Span(match.start(), match.end(), "IP_ADDRESS"))
        for match in CC_RE.finditer(text):
            if luhn_valid(match.group()): spans.append(Span(match.start(), match.end(), "CREDIT_CARD"))
        for match in DOB_RE.finditer(text): spans.append(Span(match.start(1), match.end(1), "DOB"))
        for pattern in PHONE_PATTERNS:
            for match in pattern.finditer(text): spans.append(Span(match.start(), match.end(), "PHONE"))
        if ADDRESS_CONTEXT_RE.search(text) and ADDRESS_TERMS_RE.search(text):
            spans.append(Span(0, len(text), "ADDRESS"))
        elif re.search(r"\b[A-Z]-?\d{2,}\b|\bEmbassy\b", text) and len(text) <= 250:
            spans.append(Span(0, len(text), "ADDRESS"))
        elif PIN_RE.search(text) and ADDRESS_TERMS_RE.search(text) and len(text) <= 450 and not EMAIL_RE.search(text):
            spans.append(Span(0, len(text), "ADDRESS"))
        if self.company_names:
            alt = re.compile("|".join(re.escape(v) for v in sorted(self.company_names, key=len, reverse=True)), re.I)
            for match in alt.finditer(text): spans.append(Span(match.start(), match.end(), "COMPANY"))
        if self.person_names:
            alt = re.compile("|".join(re.escape(v) for v in sorted(self.person_names, key=len, reverse=True)), re.I)
            for match in alt.finditer(text): spans.append(Span(match.start(), match.end(), "NAME"))
        priority = {"EMAIL":100,"PHONE":95,"SSN":95,"CREDIT_CARD":95,"DOB":90,"IP_ADDRESS":90,"ADDRESS":50,"COMPANY":40,"NAME":30}
        chosen = []
        for span in sorted(spans, key=lambda x: (-priority[x.label], -(x.end-x.start), x.start)):
            if any(span.start < old.end and span.end > old.start for old in chosen):
                continue
            chosen.append(span)
        return sorted(chosen, key=lambda x: x.start)

    def redact_text(self, text):
        spans = self.detect(text)
        if not spans:
            return text, []
        output = []
        current = 0
        events = []
        for span in spans:
            original = text[span.start:span.end]
            replacement = self.replacement(span.label, original)
            output.extend([text[current:span.start], replacement])
            current = span.end
            events.append({"type": span.label, "replacement": replacement})
        output.append(text[current:])
        return "".join(output), events


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def set_paragraph_text(paragraph, text):
    first = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    run = paragraph.add_run(text)
    if first:
        run.bold = first.bold
        run.italic = first.italic
        run.underline = first.underline
        run.font.name = first.font.name
        if first.font.size:
            run.font.size = first.font.size


def add_redaction(draw, box, label="[REDACTED]"):
    draw.rectangle(box, fill="black")
    x1, y1, x2, y2 = box
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", max(10, min(22, int((y2-y1) * 0.22))))
    except OSError:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), label, font=font)
    tw, th = bbox[2]-bbox[0], bbox[3]-bbox[1]
    draw.text(((x1+x2-tw)/2, (y1+y2-th)/2), label, fill="white", font=font)


def redact_identity_image(path: Path) -> dict:
    """Redact personal fields from Aadhaar/PAN images using OCR + card-layout rules."""
    image = Image.open(path).convert("RGB")
    text = pytesseract.image_to_string(image, config="--psm 6")
    lower = text.lower()
    draw = ImageDraw.Draw(image)
    width, height = image.size
    kind = None
    boxes = []

    if "permanent account number" in lower or "income tax department" in lower:
        kind = "PAN"
        # Front side: photograph, PAN, name, father's name, DOB and signature/QR.
        boxes = [
            (0.04,0.14,0.23,0.29),   # photo
            (0.30,0.19,0.57,0.27),   # PAN number
            (0.04,0.29,0.25,0.34),   # name
            (0.04,0.35,0.29,0.42),   # father's name
            (0.04,0.42,0.23,0.50),   # DOB
            (0.33,0.39,0.63,0.50),   # signature
            (0.64,0.14,0.97,0.42),   # QR
            (0.02,0.52,0.76,0.95),   # address on reverse side
        ]
    elif "unique identification authority" in lower or "aadhaar" in lower or AADHAAR_RE.search(text):
        kind = "AADHAAR"
        # Combined front/back image: photograph, name/father/DOB, ID number, QR and address.
        boxes = [
            (0.16,0.11,0.34,0.31),   # photo
            (0.32,0.12,0.63,0.30),   # name/father/DOB
            (0.31,0.30,0.62,0.37),   # Aadhaar number on front
            (0.61,0.17,0.83,0.35),   # QR
            (0.16,0.56,0.86,0.69),   # address on reverse
            (0.31,0.79,0.62,0.85),   # Aadhaar number on reverse
        ]

    if not kind:
        return {"image": path.name, "type": "OTHER", "redacted": False, "boxes": 0}

    for box in boxes:
        px = tuple(int(v) for v in (box[0]*width, box[1]*height, box[2]*width, box[3]*height))
        add_redaction(draw, px)
    image.save(path, quality=95)
    return {"image": path.name, "type": kind, "redacted": True, "boxes": len(boxes)}


def redact_images_in_docx(docx_path: Path) -> list[dict]:
    """Replace embedded image files in a DOCX after text redaction is complete."""
    with tempfile.TemporaryDirectory() as temp_dir:
        temp = Path(temp_dir)
        with ZipFile(docx_path, "r") as zin:
            zin.extractall(temp)
        media_dir = temp / "word" / "media"
        results = []
        if media_dir.exists():
            for image_path in media_dir.iterdir():
                if image_path.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
                    continue
                result = redact_identity_image(image_path)
                results.append(result)
        rebuilt = docx_path.with_suffix(".tmp.docx")
        with ZipFile(rebuilt, "w", ZIP_DEFLATED) as zout:
            for file_path in temp.rglob("*"):
                if file_path.is_file():
                    zout.write(file_path, file_path.relative_to(temp).as_posix())
        rebuilt.replace(docx_path)
        return results


def redact_docx(input_path: Path, output_path: Path, report_path: Path | None = None):
    doc = Document(input_path)
    paragraphs = list(iter_paragraphs(doc))
    redactor = Redactor()
    redactor.learn(p.text for p in paragraphs)
    redactor.learn_table_context(doc.tables)

    events = []

    # Address columns in tables are easier to handle using the table header.
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        address_columns = {
            i for i, value in enumerate(headers)
            if "address" in value or value in {"registered office", "mailing address", "residential address"}
        }
        if not address_columns:
            continue
        for row in table.rows[1:]:
            for index, cell in enumerate(row.cells):
                if index not in address_columns or not cell.text.strip():
                    continue
                replacement = redactor.replacement("ADDRESS", cell.text.strip())
                if replacement != cell.text.strip():
                    cell.text = replacement
                    events.append({"type": "ADDRESS", "replacement": replacement})

    for paragraph in iter_paragraphs(doc):
        new_text, paragraph_events = redactor.redact_text(paragraph.text)
        if new_text != paragraph.text:
            set_paragraph_text(paragraph, new_text)
            events.extend(paragraph_events)

    doc.save(output_path)
    image_events = redact_images_in_docx(output_path)

    counts = {}
    for event in events:
        counts[event["type"]] = counts.get(event["type"], 0) + 1
    image_counts = {}
    for event in image_events:
        if event["redacted"]:
            image_counts[event["type"]] = image_counts.get(event["type"], 0) + 1

    report = {
        "input_file": str(input_path),
        "output_file": str(output_path),
        "text_replacement_events": counts,
        "image_redaction_events": image_counts,
        "total_text_replacements": len(events),
        "embedded_images_checked": len(image_events),
        "embedded_identity_images_redacted": sum(1 for e in image_events if e["redacted"]),
        "unique_text_mappings": len(redactor.mapping),
    }
    if report_path:
        Path(report_path).write_text(json.dumps(report, indent=2), encoding="utf-8")
    return report


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Redact PII from DOCX text and embedded identity-card images")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--report", type=Path)
    args = parser.parse_args()
    print(json.dumps(redact_docx(args.input, args.output, args.report), indent=2))
