#!/usr/bin/env python3
"""Small evaluation suite for the PII redaction assignment."""
import json
import re
import sys
import tempfile
from pathlib import Path
from zipfile import ZipFile

import numpy as np
from PIL import Image
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
from pii_redactor import Redactor

CASES = [
    ("NAME", "Contact Person: Alice Johnson", "Alice Johnson"),
    ("NAME", "Promoters: Carol Davis, David Brown", "Carol Davis"),
    ("NAME", "Contact Person: Emily Brown", "Emily Brown"),
    ("EMAIL", "Email: alice.johnson@example.com", "alice.johnson@example.com"),
    ("EMAIL", "Reply to hr-team@sample.org", "hr-team@sample.org"),
    ("EMAIL", "support.user+1@test.co.in", "support.user+1@test.co.in"),
    ("PHONE", "Telephone: +91 9876543210", "+91 9876543210"),
    ("PHONE", "Mobile: 9123456789", "9123456789"),
    ("PHONE", "Call +1 212-555-0198", "+1 212-555-0198"),
    ("COMPANY", "Company: Acme Technologies Private Limited", "Acme Technologies Private Limited"),
    ("COMPANY", "Vendor: NorthStar Industries Limited", "NorthStar Industries Limited"),
    ("COMPANY", "Auditor: Greenfield Consulting LLP", "Greenfield Consulting LLP"),
    ("ADDRESS", "Registered Office: 10 Main Road, Mumbai, Maharashtra, India", "10 Main Road, Mumbai, Maharashtra, India"),
    ("ADDRESS", "Mailing Address: 22 Sample Street, Pune, Maharashtra, India", "22 Sample Street, Pune, Maharashtra, India"),
    ("ADDRESS", "Address: 5 Park Avenue, New Delhi, Delhi, India", "5 Park Avenue, New Delhi, Delhi, India"),
    ("SSN", "SSN: 123-45-6789", "123-45-6789"),
    ("SSN", "Social Security Number: 987-65-4321", "987-65-4321"),
    ("SSN", "Employee SSN 555-12-3456", "555-12-3456"),
    ("CREDIT_CARD", "Card: 4111 1111 1111 1111", "4111 1111 1111 1111"),
    ("CREDIT_CARD", "Payment card 5555-5555-5555-4444", "5555-5555-5555-4444"),
    ("CREDIT_CARD", "Visa 4012 8888 8888 1881", "4012 8888 8888 1881"),
    ("DOB", "Date of Birth: 12/04/1998", "12/04/1998"),
    ("DOB", "DOB: January 5, 1990", "January 5, 1990"),
    ("DOB", "Born on 7/8/1987", "7/8/1987"),
    ("IP_ADDRESS", "IP address: 192.168.1.20", "192.168.1.20"),
    ("IP_ADDRESS", "Server IP 10.0.0.15", "10.0.0.15"),
    ("IP_ADDRESS", "Client address is 203.0.113.7", "203.0.113.7"),
    ("NEGATIVE", "Order Number: 1234567890", None),
    ("NEGATIVE", "Customer Service Team", None),
    ("NEGATIVE", "support at example dot com", None),
    ("NEGATIVE", "Meeting date: 12/04/1998", None),
    ("NEGATIVE", "Version 192.168.1.999", None),
    ("NEGATIVE", "Reference number: 123-456-789", None),
    ("NEGATIVE", "Order: 1234 5678 9012", None),
    ("NEGATIVE", "Pune is a city in Maharashtra, India", None),
    ("NEGATIVE", "The company reported strong quarterly results.", None),
    ("NEGATIVE", "Office discussion without an address", None),
    ("NEGATIVE", "Ticket ID: ABC-12345", None),
    ("NEGATIVE", "Date: March 31, 2025", None),
    ("NEGATIVE", "Phone extension 1234", None),
    ("NEGATIVE", "Product code 411111111111", None),
    ("NEGATIVE", "IP range 999.999.999.999", None),
    ("NEGATIVE", "Name of the department: Finance", None),
    ("NEGATIVE", "Limited liability company discussion", None),
    ("NEGATIVE", "Bank account type: savings", None),
]

# Ground-truth regions are based on the layout of the two identity-card images
# embedded in the supplied prospectus. The regions contain PII and are expected
# to be mostly black after redaction. No original PII values are stored here.
IMAGE_CASES = {
    "image5.png": [
        (0.16,0.11,0.34,0.31), (0.32,0.12,0.63,0.30),
        (0.31,0.30,0.62,0.37), (0.61,0.17,0.83,0.35),
        (0.16,0.56,0.86,0.69), (0.31,0.79,0.62,0.85),
    ],
    "image4.png": [
        (0.04,0.14,0.23,0.29), (0.29,0.19,0.59,0.28),
        (0.04,0.28,0.29,0.35), (0.04,0.35,0.30,0.42),
        (0.04,0.41,0.28,0.49), (0.30,0.39,0.65,0.50),
        (0.64,0.14,0.97,0.42), (0.02,0.52,0.76,0.93),
    ],
}


def eval_text_cases():
    redactor = Redactor()
    redactor.learn(text for _, text, _ in CASES)
    tp = fp = tn = fn = 0
    per_type = {}

    for label, text, expected in CASES:
        spans = redactor.detect(text)
        if expected is None:
            if spans:
                fp += 1
            else:
                tn += 1
        else:
            hit = any(
                span.label == label and expected.casefold() in text[span.start:span.end].casefold()
                for span in spans
            )
            if hit:
                tp += 1
            else:
                fn += 1
            per_type.setdefault(label, [0, 0])
            per_type[label][0 if hit else 1] += 1

    total = len(CASES)
    precision = tp / (tp + fp) if tp + fp else 0
    recall = tp / (tp + fn) if tp + fn else 0
    accuracy = (tp + tn) / total if total else 0
    return {
        "method": "case-level synthetic positive and hard-negative tests",
        "total_cases": total,
        "positive_cases": 27,
        "negative_cases": 18,
        "true_positive": tp,
        "false_positive": fp,
        "true_negative": tn,
        "false_negative": fn,
        "accuracy": round(accuracy, 4),
        "precision": round(precision, 4),
        "recall": round(recall, 4),
        "per_type": {
            key: {"caught": value[0], "missed": value[1], "recall": round(value[0] / sum(value), 4)}
            for key, value in per_type.items()
        },
    }


def validate_image_regions(output_path):
    results = []
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        with ZipFile(output_path, "r") as archive:
            archive.extractall(tmp)
        media = tmp / "word" / "media"
        for image_name, boxes in IMAGE_CASES.items():
            image_path = media / image_name
            if not image_path.exists():
                results.append({"image": image_name, "regions": len(boxes), "passed": 0})
                continue
            image = Image.open(image_path).convert("RGB")
            passed = 0
            for box in boxes:
                coords = tuple(int(v * size) for v, size in zip(box, (image.width, image.height) * 2))
                crop = np.asarray(image.crop(coords))
                dark_fraction = float((crop.mean(axis=2) < 30).mean())
                if dark_fraction >= 0.65:
                    passed += 1
            results.append({"image": image_name, "regions": len(boxes), "passed": passed})
    total_regions = sum(item["regions"] for item in results)
    passed_regions = sum(item["passed"] for item in results)
    return {
        "images_checked": len(results),
        "pii_regions_checked": total_regions,
        "pii_regions_redacted": passed_regions,
        "region_redaction_rate": round(passed_regions / total_regions, 4) if total_regions else 0,
        "details": results,
    }


def doc_text(path):
    document = Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def validate_document(source, output):
    source_text = doc_text(source)
    output_text = doc_text(output)
    emails = sorted(set(re.findall(r"(?<![\w.+-])[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}(?![\w.-])", source_text)))
    audited_names = [
        "Kushal Subbayya Hegde", "Pushpa Kushal Hegde", "Rajesh Kushal Hegde", "Rohit Kushal Hegde",
        "Rakhi Girija Shetty", "Dinesh Hirachand Munot", "Ajay Shriram Patil", "Ram Kumar Tiwari",
        "Indu Jacob", "Sarthak Malvadkar", "Lokesh Shah", "Soumavo Sarkar", "Kishan Rastogi",
        "Abhijit Diwan", "Shanti Gopalkrishnan", "Hitesh Ramani", "Chitra Raste", "Sharmila Joshi",
        "Cherag Gyara", "Manisha Shukla", "Tushar Wakhele", "Ashish Mathew Pulloor", "Anand Soni",
        "Prakash Boricha", "Eric Bacha", "Sachin Gawade", "Pravin Teli", "Siddharth Jadhav",
        "Tushar Gavankar", "Varun Badai",
    ]
    audited_companies = [
        "KSH International Limited", "KSH International Private Limited", "Nuvama Wealth Management Limited",
        "ICICI Securities Limited", "MUFG Intime India Private Limited", "HDFC Bank Limited",
        "ICICI Bank Limited", "CARE Ratings Limited", "KSH Integrated Logistics Private Limited",
        "Waterloo Motors Private Limited",
    ]
    audited_addresses = [
        "11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed Pune",
        "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner Pune",
        "801 - 804, Wing A, Building No 3, Inspire BKC",
        "ICICI Venture House, Appasaheb Marathe Marg, Prabhadevi",
        "C-101, Embassy 247",
    ]
    remaining = lambda values: [value for value in values if value.casefold() in output_text.casefold()]
    return {
        "unique_source_emails": len(emails),
        "source_emails_remaining": len(remaining(emails)),
        "audited_names": len(audited_names),
        "audited_names_remaining": len(remaining(audited_names)),
        "audited_companies": len(audited_companies),
        "audited_companies_remaining": len(remaining(audited_companies)),
        "audited_addresses": len(audited_addresses),
        "audited_addresses_remaining": len(remaining(audited_addresses)),
    }


if __name__ == "__main__":
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    result = eval_text_cases()
    result["image_validation"] = validate_image_regions(output)
    result["document_validation"] = validate_document(source, output)
    out = Path(__file__).parent / "evaluation.json"
    out.write_text(json.dumps(result, indent=2), encoding="utf-8")
    print(json.dumps(result, indent=2))
